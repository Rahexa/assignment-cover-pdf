from flask import Flask, render_template, request, send_file, url_for
from weasyprint import HTML
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader, PdfWriter
import tempfile
import os
import re
import base64

# OOXML namespaces for DOCX
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
}

# Cache for base64-encoded logo so we don't read/encode it on every request
_LOGO_DATA_URI = None


def _get_run_formatting(run_el):
    """Extract formatting from a run element. Returns dict with CSS-style properties."""
    fmt = {}
    rpr = run_el.find('{%s}rPr' % NS['w'])
    if rpr is None:
        return fmt
    
    # Bold
    if rpr.find('{%s}b' % NS['w']) is not None:
        fmt['font-weight'] = 'bold'
    # Italic
    if rpr.find('{%s}i' % NS['w']) is not None:
        fmt['font-style'] = 'italic'
    # Underline
    u = rpr.find('{%s}u' % NS['w'])
    if u is not None:
        val = u.get('{%s}val' % NS['w'])
        if val not in ['none', 'false']:
            fmt['text-decoration'] = 'underline'
    
    # Font size (in half-points, convert to pt)
    sz = rpr.find('{%s}sz' % NS['w'])
    if sz is not None:
        sz_val = sz.get('{%s}val' % NS['w'])
        if sz_val:
            try:
                pt_size = int(sz_val) / 2.0
                fmt['font-size'] = f'{pt_size:.1f}pt'
            except:
                pass
    
    # Font name
    rfonts = rpr.find('{%s}rFonts' % NS['w'])
    if rfonts is not None:
        font_name = rfonts.get('{%s}ascii' % NS['w']) or rfonts.get('{%s}hAnsi' % NS['w'])
        if font_name:
            fmt['font-family'] = f'"{font_name}", serif'
    
    # Color
    color = rpr.find('{%s}color' % NS['w'])
    if color is not None:
        color_val = color.get('{%s}val' % NS['w'])
        if color_val and color_val != 'auto' and color_val.lower() != '000000':
            # Convert hex (without #) to RGB or use hex directly
            color_val = color_val.lstrip('#')
            if len(color_val) == 6:
                try:
                    # Check if it's black (default), skip if so
                    if color_val.lower() != '000000':
                        r = int(color_val[0:2], 16)
                        g = int(color_val[2:4], 16)
                        b = int(color_val[4:6], 16)
                        fmt['color'] = f'#{color_val}'
                except:
                    pass
    
    return fmt


def _format_to_css(fmt_dict):
    """Convert formatting dict to CSS string."""
    if not fmt_dict:
        return ''
    return '; '.join(f'{k}: {v}' for k, v in fmt_dict.items())


def _get_paragraph_content_items(paragraph_el, doc):
    """Extract text (with formatting) and inline images from a paragraph in order. Returns list of ('text', str, dict) or ('image', bytes, content_type)."""
    items = []
    for run_el in paragraph_el.iterchildren('{%s}r' % NS['w']):
        # Text in this run with formatting
        texts = run_el.findall('.//{%s}t' % NS['w'])
        run_text = ''.join((t.text or '') for t in texts)
        if run_text:
            fmt = _get_run_formatting(run_el)
            items.append(('text', run_text, fmt))
        # Inline drawing (image) in this run
        blip = run_el.find('.//{%s}blip' % NS['a'])
        if blip is not None:
            r_id = blip.get('{%s}embed' % NS['r']) or blip.get('embed')
            if r_id and hasattr(doc.part, 'related_parts'):
                try:
                    image_part = doc.part.related_parts[r_id]
                    blob = getattr(image_part, 'blob', None) or getattr(image_part, '_blob', None)
                    if blob:
                        ct = getattr(image_part, 'content_type', 'image/png') or 'image/png'
                        items.append(('image', blob, ct))
                except Exception:
                    pass
    return items


def _get_paragraph_formatting(paragraph_el):
    """Extract paragraph-level formatting (alignment, spacing). Returns dict."""
    fmt = {}
    ppr = paragraph_el.find('{%s}pPr' % NS['w'])
    if ppr is None:
        return fmt
    
    # Alignment
    jc = ppr.find('{%s}jc' % NS['w'])
    if jc is not None:
        align = jc.get('{%s}val' % NS['w'])
        if align in ['left', 'center', 'right', 'both', 'justify']:
            fmt['text-align'] = align if align != 'both' else 'justify'
    
    # Spacing (before/after in twips, 1pt = 20 twips)
    spacing = ppr.find('{%s}spacing' % NS['w'])
    if spacing is not None:
        before = spacing.get('{%s}before' % NS['w'])
        after = spacing.get('{%s}after' % NS['w'])
        if before:
            try:
                pt_before = int(before) / 20.0
                fmt['margin-top'] = f'{pt_before:.1f}pt'
            except:
                pass
        if after:
            try:
                pt_after = int(after) / 20.0
                fmt['margin-bottom'] = f'{pt_after:.1f}pt'
            except:
                pass
    
    return fmt


def _element_content_to_html(element_el, doc, tag='p'):
    """Turn a block element (paragraph or cell content) into HTML, preserving formatting, text and images in order."""
    parts = []
    if element_el.tag != '{%s}p' % NS['w']:
        # Maybe a single paragraph wrapper
        for p_el in element_el.iterchildren('{%s}p' % NS['w']):
            p_fmt = _get_paragraph_formatting(p_el)
            p_style = _format_to_css(p_fmt)
            run_parts = []
            for item in _get_paragraph_content_items(p_el, doc):
                if item[0] == 'text':
                    text, fmt = item[1], item[2]
                    s = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    run_style = _format_to_css(fmt)
                    if run_style:
                        run_parts.append(f'<span style="{run_style}">{s}</span>')
                    else:
                        run_parts.append(s)
                else:
                    # image: (_, blob, content_type)
                    b64 = base64.b64encode(item[1]).decode('ascii')
                    mt = (item[2] if len(item) > 2 else 'image/png').split(';')[0].strip()
                    if mt.startswith('image/'):
                        run_parts.append(f'<img src="data:{mt};base64,{b64}" style="max-width:100%;height:auto;display:block;margin:8px 0;" />')
            inner = ''.join(run_parts)
            if inner.strip():
                if p_style:
                    parts.append(f'<{tag} style="{p_style}">{inner}</{tag}>')
                else:
                    parts.append(f'<{tag}>{inner}</{tag}>')
        return ''.join(parts)
    
    # Single paragraph element
    p_fmt = _get_paragraph_formatting(element_el)
    p_style = _format_to_css(p_fmt)
    for item in _get_paragraph_content_items(element_el, doc):
        if item[0] == 'text':
            text, fmt = item[1], item[2]
            s = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            run_style = _format_to_css(fmt)
            if run_style:
                parts.append(f'<span style="{run_style}">{s}</span>')
            else:
                parts.append(s)
        else:
            b64 = base64.b64encode(item[1]).decode('ascii')
            mt = (item[2] if len(item) > 2 else 'image/png').split(';')[0].strip()
            if mt.startswith('image/'):
                parts.append(f'<img src="data:{mt};base64,{b64}" style="max-width:100%;height:auto;display:block;margin:8px 0;" />')
    inner = ''.join(parts)
    if tag == 'td':
        return inner
    if not inner.strip():
        return ''
    if p_style:
        return f'<{tag} style="{p_style}">{inner}</{tag}>'
    return f'<{tag}>{inner}</{tag}>'


def _get_logo_data_uri():
    """Return a cached base64 data URI for the logo image."""
    global _LOGO_DATA_URI
    if _LOGO_DATA_URI is not None:
        return _LOGO_DATA_URI

    logo_path = os.path.join(app.root_path, 'static', 'images', 'puclogo.png')
    try:
        with open(logo_path, 'rb') as logo_file:
            logo_data = base64.b64encode(logo_file.read()).decode('utf-8')
            _LOGO_DATA_URI = f'data:image/png;base64,{logo_data}'
    except FileNotFoundError:
        # Fallback to HTTP URL if file not found
        _LOGO_DATA_URI = None

    return _LOGO_DATA_URI


app = Flask(__name__)


@app.route('/health')
@app.route('/ping')
def health():
    """Lightweight endpoint for Uptime Robot / cron to keep Render free tier alive. No heavy imports or work."""
    return 'ok', 200, {'Content-Type': 'text/plain'}


@app.route('/')
def index():
    return render_template("index.html")

@app.route('/generate', methods=['POST'])
def generate():
    data = request.form
    files = request.files
    template_key = data.get('template', 'template1')
    submitted_to = data.get('submitted_to', '').strip()
    submitted_designation = data.get('submitted_designation', '').strip()
    output_format = data.get('output_format', 'pdf')
    output_type = data.get('output_type', 'cover')
    cover_type = data.get('cover_type', 'assignment')
    batch = data.get('batch', '').strip()
    section = data.get('section', '').strip()
    session = data.get('session', '').strip()
    performance_date = data.get('performance_date', '').strip()

    # Build a safe filename based on student ID
    raw_id = data.get('student_id', '').strip()
    safe_id = re.sub(r'[^0-9A-Za-z_-]', '', raw_id) or 'assignment'

    if output_type == 'cover' and output_format == 'docx':
        # Generate a simple DOCX version of the cover page
        document = Document()
        heading = 'Lab Report Cover Page' if cover_type == 'lab' else 'Assignment Cover Page'
        no_label = 'Lab Report No.' if cover_type == 'lab' else 'Assignment No.'
        name_label = 'Lab Report Name' if cover_type == 'lab' else 'Assignment Name'
        document.add_heading(heading, level=0)

        document.add_paragraph(f"{no_label}: {data['assignment_no']}")
        document.add_paragraph(f"Course Code: {data['course_code']}")
        document.add_paragraph(f"Course Title: {data['course_title']}")
        document.add_paragraph(f"{name_label}: {data['assignment_name']}")
        if performance_date:
            document.add_paragraph(f"Date of Performance: {performance_date}")
        document.add_paragraph(f"Date of Submission: {data['submission_date']}")
        document.add_paragraph("")
        document.add_paragraph(f"Student Name: {data['student_name']}")
        document.add_paragraph(f"ID: {data['student_id']}")
        if batch:
            document.add_paragraph(f"Batch: {batch}")
        if section:
            document.add_paragraph(f"Section: {section}")
        if session:
            document.add_paragraph(f"Session: {session}")
        document.add_paragraph("")
        if submitted_to or submitted_designation:
            document.add_paragraph("Submitted to:")
            if submitted_to:
                document.add_paragraph(submitted_to)
            if submitted_designation:
                document.add_paragraph(submitted_designation)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as docx_file:
            document.save(docx_file.name)
            return send_file(
                docx_file.name,
                as_attachment=True,
                download_name=f"{safe_id}.docx",
            )

    # Decide which HTML template to use for PDF
    if template_key == 'template2':
        template_file = 'cover2.html'
    elif template_key == 'template3':
        template_file = 'cover3.html'
    else:
        template_file = 'cover.html'

    # Embed logo as base64 data URI (cached, no file/HTTP requests per request)
    logo_url = _get_logo_data_uri()
    if not logo_url:
        # Fallback to HTTP URL if file not found or failed to load
        logo_url = request.url_root.rstrip('/') + url_for('static', filename='images/puclogo.png')
    
    # Always render the cover HTML for PDF generation
    html = render_template(
        template_file,
        assignment_no=data['assignment_no'],
        course_code=data['course_code'],
        course_title=data['course_title'],
        assignment_name=data['assignment_name'],
        performance_date=performance_date,
        submission_date=data['submission_date'],
        student_name=data['student_name'],
        student_id=data['student_id'],
        submitted_to=submitted_to,
        submitted_designation=submitted_designation,
        cover_type=cover_type,
        batch=batch,
        section=section,
        session=session,
        logo_url=logo_url,
    )

    # Generate cover PDF (base64 logo eliminates slow HTTP/file requests)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as cover_file:
        HTML(string=html, base_url=request.url_root).write_pdf(cover_file.name)
        cover_path = cover_file.name

    # If only cover page is requested, just return the cover PDF
    if output_type == 'cover':
        return send_file(
            cover_path,
            as_attachment=True,
            download_name=f"{safe_id}.pdf",
        )

    # Otherwise, merge with uploaded assignment (PDF, DOC, DOCX)
    assignment = files.get('assignment_file')
    if not assignment or assignment.filename == '':
        return "No assignment file provided for merge.", 400

    _, ext = os.path.splitext(assignment.filename)
    ext = ext.lower()

    # Handle different file types
    if ext == '.pdf':
        # Direct PDF - save and use
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as assignment_file:
            assignment.save(assignment_file.name)
            assignment_path = assignment_file.name
    elif ext == '.docx':
        # Convert DOCX to PDF first
        try:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_file:
                assignment.save(docx_file.name)
                docx_path = docx_file.name
            
            # Read DOCX and convert to HTML (preserving formatting, images, etc.)
            doc = Document(docx_path)
            # Minimal base CSS - DOCX formatting will override via inline styles
            html_parts = ['<html><head><meta charset="UTF-8"><style>body{font-family:"Times New Roman",Times,serif;padding:20px;line-height:1.5;color:#000;}p{margin:6px 0;}table{border-collapse:collapse;width:100%;margin:10px 0;}td,th{border:1px solid #000;padding:4px 8px;vertical-align:top;}img{max-width:100%;height:auto;display:block;margin:4px 0;}span{display:inline;}</style></head><body>']

            for element in doc.element.body:
                if isinstance(element, CT_P):
                    p_html = _element_content_to_html(element, doc, 'p')
                    if p_html:
                        html_parts.append(p_html)
                elif isinstance(element, CT_Tbl):
                    tbl = Table(element, doc)
                    html_parts.append('<table>')
                    for row in tbl.rows:
                        html_parts.append('<tr>')
                        for cell in row.cells:
                            cell_parts = []
                            for p_el in cell._tc.iterchildren('{%s}p' % NS['w']):
                                cell_parts.append(_element_content_to_html(p_el, doc, 'td'))
                            cell_html = ''.join(cell_parts) or ' '
                            html_parts.append(f'<td>{cell_html}</td>')
                        html_parts.append('</tr>')
                    html_parts.append('</table>')

            html_parts.append('</body></html>')
            html_content = ''.join(html_parts)
            
            # Convert HTML to PDF
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_file:
                HTML(string=html_content, base_url=request.url_root).write_pdf(pdf_file.name)
                assignment_path = pdf_file.name
            
            # Clean up DOCX temp file
            try:
                os.unlink(docx_path)
            except:
                pass
        except Exception as e:
            return f"Error converting DOCX to PDF: {str(e)}. Please ensure the file is a valid DOCX document.", 400
    elif ext == '.doc':
        # DOC format (older format) - python-docx doesn't support it
        return "DOC files (older format) are not supported. Please convert your file to DOCX or PDF format first.", 400
    else:
        return f"Unsupported file type: {ext}. Please use PDF, DOC, or DOCX files.", 400

    writer = PdfWriter()
    # Add cover pages
    cover_reader = PdfReader(cover_path)
    for page in cover_reader.pages:
        writer.add_page(page)

    # Add assignment pages
    assignment_reader = PdfReader(assignment_path)
    for page in assignment_reader.pages:
        writer.add_page(page)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as merged_file:
        writer.write(merged_file.name)
        return send_file(
            merged_file.name,
            as_attachment=True,
            download_name=f"{safe_id}.pdf",
        )

if __name__ == '__main__':
    app.run(debug=True)
